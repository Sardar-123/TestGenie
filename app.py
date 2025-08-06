import os
from flask import Flask, render_template, request, flash, send_file, redirect
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
import pandas as pd
import json
from docx import Document
import openai
import io
import re
import zipfile
import xml.dom.minidom
import markdown
from PIL import Image
import base64
import PyPDF2
import time

load_dotenv()                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                       

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'txt', 'csv', 'xlsx', 'xls', 'docx', 'json', 'pdf', 'xml', 'md', 'html', 'zip', 'png', 'jpg', 'jpeg', 'gif'}
MAX_FILE_SIZE_MB = 50
UPLOAD_RETENTION_SECONDS = 3600  # 1 hour

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.secret_key = 'your_secret_key'

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_file_preview(filepath, ext):
    try:
        if ext == 'txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                lines = []
                for _ in range(5):
                    line = f.readline()
                    if not line:
                        break
                    lines.append(line.rstrip())
                return '\n'.join(lines) if lines else 'No content to preview.'
        elif ext == 'csv':
            df = pd.read_csv(filepath)
            return df.head().to_html(classes='table table-bordered')
        elif ext in ['xlsx', 'xls']:
            df = pd.read_excel(filepath)
            return df.head().to_html(classes='table table-bordered')
        elif ext == 'json':
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return f"<pre>{json.dumps(data, indent=2)[:500]}</pre>"
        elif ext == 'docx':
            doc = Document(filepath)
            text = '\n'.join([para.text for para in doc.paragraphs[:5]])
            return f"<pre>{text[:1000] if text else 'No text content found in document.'}</pre>"
        elif ext == 'pdf':
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                if reader.pages:
                    text = reader.pages[0].extract_text()
                    return f"<pre>{text[:1000] if text else 'No text content found on first page.'}</pre>"
                else:
                    return "No pages found in PDF."
        elif ext == 'xml':
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
                dom = xml.dom.minidom.parseString(content)
                pretty_xml = dom.toprettyxml()
                return f"<pre>{pretty_xml[:1000]}</pre>"
        elif ext == 'md':
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
                html = markdown.markdown(content)
            return f"<div style='background:#f8f9fa;padding:1em;'>{html[:1000]}</div>"
        elif ext == 'html':
            with open(filepath, 'r', encoding='utf-8') as f:
                html = f.read()
            return f"<pre>{html[:1000]}</pre>"
        elif ext in ['png', 'jpg', 'jpeg', 'gif']:
            try:
                with open(filepath, 'rb') as f:
                    img = Image.open(f)
                    buffered = io.BytesIO()
                    # Convert to RGB if necessary (for PNG with transparency)
                    if img.mode in ('RGBA', 'LA'):
                        img = img.convert('RGB')
                    img.save(buffered, format='JPEG')
                    img_str = base64.b64encode(buffered.getvalue()).decode()
                    return f"<img src='data:image/jpeg;base64,{img_str}' style='max-width:300px;max-height:300px;'/>"
            except Exception as img_error:
                return f"Image preview unavailable: {str(img_error)}"
        elif ext == 'zip':
            with zipfile.ZipFile(filepath, 'r') as zipf:
                file_list = zipf.namelist()
            return f"<pre>ZIP Contents ({len(file_list)} files):\n" + '\n'.join(file_list[:10]) + ("\n... and more" if len(file_list) > 10 else "") + "</pre>"
        else:
            return "Preview not available for this file type."
    except Exception as e:
        return f"Error previewing file: {str(e)}"

def generate_ai_test_cases(filepath, test_type, test_level, industry, output_format, num_cases, code_language, strict_scenarios, ext):
    try:
        # Handle different file types for content extraction
        if ext == 'txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read(1000)
        elif ext == 'csv':
            df = pd.read_csv(filepath)
            content = df.head(10).to_string()
        elif ext in ['xlsx', 'xls']:
            df = pd.read_excel(filepath)
            content = df.head(10).to_string()
        elif ext == 'json':
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
            content = json.dumps(data, indent=2)[:1000]
        elif ext == 'docx':
            doc = Document(filepath)
            content = '\n'.join([para.text for para in doc.paragraphs[:10]])
        elif ext == 'pdf':
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                if reader.pages:
                    content = reader.pages[0].extract_text()[:1000]
                else:
                    content = "PDF document with no extractable text"
        elif ext == 'xml':
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read(1000)
        elif ext == 'md':
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read(1000)
        elif ext == 'html':
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read(1000)
        elif ext in ['png', 'jpg', 'jpeg', 'gif']:
            content = f"Image file: {filepath}. Image analysis would require computer vision capabilities."
        elif ext == 'zip':
            with zipfile.ZipFile(filepath, 'r') as zipf:
                file_list = zipf.namelist()
            content = f"ZIP archive containing: {', '.join(file_list[:10])}"
        else:
            content = "File type not supported for content analysis."
    except Exception as e:
        content = f"Could not extract content from file: {str(e)}"
    # Dynamic prompt with explicit instructions and examples
    strict_note = "Generate code only for the scenario(s) described below. Do not add extra scenarios." if strict_scenarios else ""
    format_instructions = {
        'Manual': (
            f'{strict_note} Generate {num_cases} detailed manual test cases as step-by-step instructions. Output only the test cases, no explanation.'
        ),
        'Gherkin': (
            f'{strict_note} Generate {num_cases} test cases in Gherkin syntax (Given-When-Then). Output only valid Gherkin feature file content, no explanation.\n'
            'Example:\nFeature: User login\n  Scenario: Successful login\n    Given the user is on the login page\n    When the user enters valid credentials\n    Then the user should be redirected to the dashboard'
        ),
        'API (Code)': (
            f'{strict_note} Generate {num_cases} API test cases as code in {code_language}. Use popular libraries for API testing (e.g., axios/chai for JavaScript, requests/pytest for Python, RestAssured/JUnit for Java). Output only valid {code_language} code, no explanation.\n'
            'Example (JavaScript):\nconst axios = require("axios");\nconst { expect } = require("chai");\ndescribe("Login API", () => {\n  it("should return 200 and a token for valid credentials", async () => {\n    const response = await axios.post("http://your-api-url/api/login", { username: "user", password: "pass" });\n    expect(response.status).to.equal(200);\n    expect(response.data).to.have.property("token");\n  });\n});'
        ),
        'Selenium': (
            f'{strict_note} Generate {num_cases} Selenium test scripts for automation in {code_language}. Output only valid {code_language} code using Selenium, no explanation.\n'
            'Example (Python):\nfrom selenium import webdriver\nfrom selenium.webdriver.common.by import By\n\ndef test_login():\n    driver = webdriver.Chrome()\n    driver.get("https://example.com/login")\n    driver.find_element(By.ID, "username").send_keys("user")\n    driver.find_element(By.ID, "password").send_keys("pass")\n    driver.find_element(By.ID, "login").click()\n    assert "Dashboard" in driver.title\n    driver.quit()'
        ),
        'Cypress': (
            f'{strict_note} Generate {num_cases} Cypress test scripts for automation in {code_language}. Output only valid {code_language} code using Cypress, no explanation.\n'
            'Example (JavaScript):\ndescribe("Login", () => {\n  it("should login successfully", () => {\n    cy.visit("/login");\n    cy.get("#username").type("user");\n    cy.get("#password").type("pass");\n    cy.get("#login").click();\n    cy.contains("Dashboard");\n  });\n});'
        ),
        'Playwright': (
            f'{strict_note} Generate {num_cases} Playwright test scripts for automation in {code_language}. Output only valid {code_language} code using Playwright, no explanation.\n'
            'Example (JavaScript):\nconst { test, expect } = require("@playwright/test");\ntest("login", async ({ page }) => {\n  await page.goto("/login");\n  await page.fill("#username", "user");\n  await page.fill("#password", "pass");\n  await page.click("#login");\n  await expect(page).toHaveText("Dashboard");\n});'
        )
    }
    prompt = (
        f"You are an expert software tester. If the provided content does not match the selected testing type, respond with: 'The provided content does not match the selected testing type.'\n"
        f"Testing Type: {test_type}\nTesting Level: {test_level}\nIndustry: {industry}\nOutput Format: {output_format}\nLanguage: {code_language}\n"
        f"File Content: {content}\n"
        f"{format_instructions.get(output_format, '')} "
        f"Generate {num_cases} {test_level} test cases for {test_type} testing in {industry} industry based on the above requirements."
    )
    try:
        client = openai.AzureOpenAI(
            api_key=os.getenv("AZURE_OPENAI_API_KEY"),
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
            api_version="2023-03-15-preview"
        )
        response = client.chat.completions.create(
            model=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1000
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error generating test cases: {e}"

generated_cases_cache = {}

def cleanup_uploads():
    now = time.time()
    for fname in os.listdir(UPLOAD_FOLDER):
        fpath = os.path.join(UPLOAD_FOLDER, fname)
        if os.path.isfile(fpath):
            if now - os.path.getmtime(fpath) > UPLOAD_RETENTION_SECONDS:
                os.remove(fpath)

def is_safe_file_content(filepath, ext):
    try:
        # For binary file types, we don't need to check content safety
        if ext in ['png', 'jpg', 'jpeg', 'gif', 'pdf', 'xlsx', 'xls', 'docx', 'zip']:
            return True
        
        # For text-based files, check content safety
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read(10000)  # Read first 10KB for safety check
        
        # Basic check: no binary, no dangerous HTML/script tags
        if '\x00' in content:
            return False
        dangerous = ['<script', '<iframe', '<object', '<embed', 'eval(', 'base64,']
        for d in dangerous:
            if d in content.lower():
                return False
        return True
    except Exception:
        # If we can't read the file, assume it's safe for binary types
        if ext in ['png', 'jpg', 'jpeg', 'gif', 'pdf', 'xlsx', 'xls', 'docx', 'zip']:
            return True
        return False

@app.route('/', methods=['GET', 'POST'])
def index():
    cleanup_uploads()
    test_types = ['Database', 'API', 'Mobile', 'Manual', 'Performance']
    test_levels = ['Unit', 'Integration', 'System', 'Acceptance', 'Smoke', 'Sanity', 'Regression']
    industries = ['Financial/Banking', 'Healthcare', 'E-commerce']
    output_formats = ['Manual', 'Gherkin', 'Selenium', 'Cypress', 'Playwright']
    preview = None
    generated_cases = None
    if request.method == 'POST':
        file = request.files.get('file')
        # File size validation
        file.seek(0, os.SEEK_END)
        size_mb = file.tell() / (1024 * 1024)
        file.seek(0)
        if size_mb > MAX_FILE_SIZE_MB:
            flash(f'File too large (>{MAX_FILE_SIZE_MB}MB).')
            return render_template('index.html',
                test_types=test_types,
                test_levels=test_levels,
                industries=industries,
                output_formats=output_formats,
                preview=preview,
                generated_cases=generated_cases)
        # ...existing code...
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            print(f"File saved to {filepath}")
            ext = filename.rsplit('.', 1)[1].lower()
            
            # Check file safety based on file type
            if not is_safe_file_content(filepath, ext):
                os.remove(filepath)
                flash('File content is not safe or valid.')
                return render_template('index.html',
                    test_types=test_types,
                    test_levels=test_levels,
                    industries=industries,
                    output_formats=output_formats,
                    preview=preview,
                    generated_cases=generated_cases)
            
            # Generate file preview
            preview = get_file_preview(filepath, ext)
            test_type = request.form.get('test_type')
            test_level = request.form.get('test_level')
            industry = request.form.get('industry')
            output_format = request.form.get('output_format')
            generate_ai = request.form.get('generate_ai')
            num_cases = int(request.form.get('num_cases', 5))
            code_language = request.form.get('code_language', 'Python')
            strict_scenarios = bool(request.form.get('strict_scenarios'))
            if generate_ai:
                generated_cases = generate_ai_test_cases(filepath, test_type, test_level, industry, output_format, num_cases, code_language, strict_scenarios, ext)
                if not generated_cases or not generated_cases.strip():
                    generated_cases = 'No test cases could be generated for the provided scenario(s). Please check your input or try again with less strict settings.'
                generated_cases_cache['content'] = generated_cases
        else:
            flash('Invalid file type or no file selected.')
    return render_template('index.html',
                           test_types=test_types,
                           test_levels=test_levels,
                           industries=industries,
                           output_formats=output_formats,
                           preview=preview,
                           generated_cases=generated_cases)

@app.route('/download_test_cases')
def download_test_cases():
    content = generated_cases_cache.get('content', '')
    if not content:
        flash('No test cases to download.')
        return redirect('/')
    # Clean up content: remove any HTML tags if present, ensure plain text
    content = re.sub(r'<[^>]+>', '', content)
    content = content.replace('\r\n', '\n').replace('\r', '\n')
    return send_file(
        io.BytesIO(content.encode('utf-8')),
        mimetype='text/plain',
        as_attachment=True,
        download_name='test_cases.txt'
    )

if __name__ == '__main__':
    app.run(debug=True)
